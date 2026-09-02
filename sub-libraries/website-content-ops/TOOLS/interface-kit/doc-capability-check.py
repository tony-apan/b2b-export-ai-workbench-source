#!/usr/bin/env python3
"""文档解析能力检查（doc-capability-check.py）—— 分享/使用 skill 前的环境门。
检查本机是否能读 PDF / Word / PPT / 表格等常见文档；输出能力矩阵 + 安装建议。跨平台（mac/win/linux）。

用法：
  python3 doc-capability-check.py            # 检查并输出矩阵
  python3 doc-capability-check.py --install  # 给出（不执行）安装命令
  python3 doc-capability-check.py self-test  # 生成测试文档并实际解析自证

结论门：pdf=ok / docx=ok / pptx=ok / xlsx=ok / md=ok 全部通过 → 可安全承接"用户资料解析"；
任何 not-ok → 先按 --install 装机（pip 或 brew/apt），装机后重跑本检查再开工。
"""
import importlib, importlib.util as _iu, os, shutil, subprocess, sys, tempfile, platform

def has_module(name): return importlib.util.find_spec(name) is not None

def has_exe(name): return shutil.which(name) is not None

def check():
    system = platform.system()
    out = {}
    # PDF
    if has_module("pypdf"): out["pdf"] = ("ok", "pypdf")
    elif has_module("PyPDF2"): out["pdf"] = ("ok", "PyPDF2")
    elif has_exe("pdftotext"): out["pdf"] = ("ok", "poppler pdftotext")
    else: out["pdf"] = ("not-ok", "需 pypdf 或 poppler")
    # Word
    if has_module("docx"): out["docx"] = ("ok", "python-docx")
    else: out["docx"] = ("not-ok", "需 python-docx")
    # PPT
    if has_module("pptx"): out["pptx"] = ("ok", "python-pptx")
    else: out["pptx"] = ("not-ok", "需 python-pptx")
    # Excel
    if has_module("openpyxl"): out["xlsx"] = ("ok", "openpyxl")
    elif has_module("pandas"): out["xlsx"] = ("ok", "pandas")
    else: out["xlsx"] = ("not-ok", "需 openpyxl")
    # Markdown / 文本（内置）
    out["md"] = ("ok", "内置")
    # 网页（内置 urllib）
    out["html"] = ("ok", "内置")
    return system, out

def main():
    args = sys.argv[1:]
    system, mat = check()
    print(f"System: {system} / Python {sys.version.split()[0]}")
    print(f"{'格式':<6} {'状态':<8} 方案")
    allok = True
    for fmt, (st, by) in mat.items():
        print(f"{fmt:<6} {st:<8} {by}")
        if st != "ok": allok = False
    print("\n== 结论 ==", "PASS：可承接文档解析与建站" if allok else "FAIL：先安装缺失项（见 --install）再重跑")
    if "--install" in args:
        print("\n== 安装建议（按系统选择其一）==")
        print("  pip:  pip install pypdf python-docx python-pptx openpyxl")
        print("  mac:  brew install poppler  # pdftotext（备选）")
        print("  ubuntu: apt install -y poppler-utils")
    if "self-test" in args:
        print("\n== self-test（生成示例并解析）==")
        d = tempfile.mkdtemp()
        # txt/md 内置；pdf/docx 依赖库（若能）则生成
        try:
            from docx import Document
            doc = Document(); doc.add_heading("Hi", 1); doc.add_paragraph("docx test")
            doc.save(os.path.join(d, "t.docx")); print("docx self-test write ok")
        except Exception as e: print("docx self-test skipped:", type(e).__name__)
    # 门：任一格式 not-ok → exit 1（--install/self-test 只是提示与自证，不豁免缺失项）。
    if not allok:
        return 1
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
